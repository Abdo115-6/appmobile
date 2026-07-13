#!/usr/bin/env python3
"""
Generate a 6-week internship report (Rapport de Stage) for Abdelkarim Erraji
at Mallzellij, supervised by Driss Rokh, from L'EMSI.
Based on actual development notes (NOTES.docx) — 17 challenges.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────────

def setup_page(doc, size="A4"):
    section = doc.sections[0]
    if size == "A4":
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.top_margin = section.bottom_margin = Cm(2.54)
        section.left_margin = section.right_margin = Cm(3.18)
    else:
        section.page_width, section.page_height = Inches(8.5), Inches(11.0)
        section.top_margin = section.bottom_margin = Inches(1.0)
        section.left_margin = section.right_margin = Inches(1.25)
    section.orientation = WD_ORIENTATION.PORTRAIT


def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)

    for n, size, color in [
        (1, 18, RGBColor(0x1A, 0x3C, 0x6E)),
        (2, 14, RGBColor(0x1A, 0x3C, 0x6E)),
        (3, 12, RGBColor(0x2C, 0x5F, 0x8A)),
    ]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)

    title = doc.styles["Title"]
    title.font.name = "Calibri Light"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = "Mettre a jour le sommaire (clic droit -> Mettre a jour le champ)"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    for x in (fldChar1, instrText, fldChar2, fldChar3, fldChar4):
        run._r.append(x)


def shade_paragraph(paragraph, hex_color="F2F4F7"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def add_table(doc, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(header):
        hdr_cells[i].text = col_name
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = str(value)
    return table


def add_bullet_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


# ── Cover page ───────────────────────────────────────────────────────────────

def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ECOLE MAROCAINE DES SCIENCES DE L'INGENIEUR")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L'EMSI")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    p = doc.add_paragraph("RAPPORT DE STAGE", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Application Android de Gestion d'Inventaire")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("chez Mallzellij")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Presente par : Abdelkarim ERRaji")
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Encadrant : Driss ROKH")
    run.font.size = Pt(12)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Annee universitaire 2025 - 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ── Header / Footer ─────────────────────────────────────────────────────────

def setup_header_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)
    header = section.header.paragraphs[0]
    header.text = "Rapport de Stage - Mallzellij"
    header.style = doc.styles["Header"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ── Content ──────────────────────────────────────────────────────────────────

def add_general_introduction(doc):
    doc.add_paragraph("Introduction Generale", style="Heading 1")
    doc.add_paragraph(
        "Ce rapport presente le stage de fin d'etudes realise durant six semaines "
        "au sein de l'entreprise Mallzellij, dans le cadre de la formation en "
        "genie logiciel a l'Ecole Marocaine des Sciences de l'Ingenieur (L'EMSI). "
        "L'objectif principal de ce stage etait de concevoir et developper une "
        "application mobile Android permettant la gestion efficace de l'inventaire "
        "des articles du magasin."
    )
    doc.add_paragraph(
        "Le projet a suivi un developpement iteratif et incrementel, decompose en "
        "17 defis techniques releves successivement, couvrant la mise en place de "
        "l'environnement de developpement, la creation de l'interface utilisateur, "
        "l'integration avec la base de donnees existante de l'entreprise (Sage X3), "
        "le deploiement sur serveur, la gestion des roles, la fonctionnalite de "
        "scan de codes QR, et l'ajout du formulaire d'inventaire complet."
    )
    doc.add_paragraph(
        "Ce document s'articule autour de plusieurs parties : la presentation de "
        "l'entreprise daccueil, lanalyse des besoins et la conception de la "
        "solution, larchitecture technique du systeme, les details "
        "dimplementation suivant les 17 defis releves, ainsi quun bilan "
        "personnel et professionnel du stage."
    )


def add_presentation_entreprise(doc):
    doc.add_paragraph("Presentation de l'Entreprise", style="Heading 1")

    doc.add_paragraph("1.1 Presentation generale", style="Heading 2")
    doc.add_paragraph(
        "Mallzellij est une entreprise specialisee dans la distribution et la "
        "gestion de produits. L'entreprise gere un vaste catalogue d'articles "
        "repportis sur plusieurs sites (FBC, FMD), et necessite un outil efficace "
        "pour suivre les stocks et realiser des inventaires en temps reel. "
        "L'entreprise utilise le systeme d'information Sage X3 pour la gestion "
        "de sa base de donnees SQL Server."
    )

    doc.add_paragraph("1.2 Encadrement et contexte", style="Heading 2")
    doc.add_paragraph(
        "Le stage s'est deroule sous la direction de Driss Rokh, en tant "
        "qu'encadrant technique. L'equipe de developpement travaille en "
        "collaboration etroite avec le service logistique pour repondre aux "
        "besoins operationnels de l'entreprise."
    )

    doc.add_paragraph("1.3 Mission du stagiaire", style="Heading 2")
    add_bullet_list(doc, [
        "Choisir les technologies adaptees (SQL Server, Spring Boot, Android Studio)",
        "Connecter l'application a la base de donnees existante de l'entreprise",
        "Creer l'interface utilisateur avec authentification",
        "Afficher les articles et les stocks par site",
        "Integrer le scan de codes QR pour identifier les articles",
        "Creer un formulaire d'inventaire complet",
        "Gerer les droits d'acces par role (admin/utilisateur)",
        "Deployer la solution sur serveur pour un fonctionnement 24h/24",
    ])


def add_conception(doc):
    doc.add_paragraph("Analyse des Besoins et Conception", style="Heading 1")

    doc.add_paragraph("2.1 Besoins fonctionnels", style="Heading 2")
    doc.add_paragraph(
        "L'analyse des besoins a revele les exigences fonctionnelles suivantes, "
        "derivees des 17 defis releves pendant le stage :"
    )
    add_numbered_list(doc, [
        "Authentification securisee des utilisateurs (email + mot de passe)",
        "Affichage de la liste des articles disponibles dans le catalogue",
        "Affichage du stock detaille par site (FBC, FMD) apres selection d'un article",
        "Recherche d'articles par nom en temps reel (debounce 400ms)",
        "Scan de codes QR pour identifier un article (reference ITMREF_0)",
        "Formulaire d'inventaire : equipe, depot, zone, article, palettes, cartons, metres carres",
        "Dropdowns fixes (equipe, depot, zone) qui persistent apres rechargement",
        "Dialog de confirmation avant insertion en base avec recapitulatif",
        "Gestion des roles : admin (articles + inventaire) et user1 (articles uniquement)",
        "Navigation via sidebar entre articles et inventaire",
        "Affichage du prix des articles",
    ])

    doc.add_paragraph("2.2 Besoins non fonctionnels", style="Heading 2")
    add_bullet_list(doc, [
        "Interface utilisateur intuitive et ergonomique (Material Design)",
        "Compatibilite Android API 24+ (Android 7.0 et au-dela)",
        "Securite des donnees sensibles (mots de passe)",
        "Deploiement sur serveur VMware pour fonctionnement 24h/24",
        "Reseau local uniquement (securite : pas d'acces externe)",
        "Synchronisation avec les donnees de l'entreprise (Sage X3)",
    ])

    doc.add_paragraph("2.3 Modele de donnees", style="Heading 2")
    doc.add_paragraph(
        "Le systeme s'appuie sur les tables suivantes, mappees aux tables "
        "de la base de donnees Sage X3 existante :"
    )
    add_table(doc,
        ["Entite", "Table SQL Server", "Description"],
        [
            ["Utilisateur", "YMOBILE_USER", "Comptes utilisateurs avec roles (admin/user1)"],
            ["Article", "ITM_MASTER", "Catalogue des articles (ITMREF_0, ITMDES1_0)"],
            ["Mouvement stock", "ITM_MVT", "Mouvements de stock par site (PHYALL_0)"],
            ["Inventaire", "YINV", "Enregistrements d'inventaire (nouveau)"],
            ["Prix", "SPRC_LIST", "Listes de prix par article (liste T11)"],
        ]
    )
    doc.add_paragraph()


def add_architecture(doc):
    doc.add_paragraph("Architecture Technique", style="Heading 1")

    doc.add_paragraph("3.1 Vue d'ensemble", style="Heading 2")
    doc.add_paragraph(
        "L'application suit une architecture client-serveur. Le frontend Android "
        "communique avec le backend Spring Boot via une API RESTful, qui interagit "
        "avec la base de donnees SQL Server existante de l'entreprise Sage X3. "
        "Le serveur backend est deploie sur une VM VMware pour un fonctionnement "
        "continu (24h/24)."
    )

    doc.add_paragraph("3.2 Backend - Spring Boot", style="Heading 2")
    doc.add_paragraph(
        "Le backend est developpe en Java 17 avec Spring Boot 3.4.4. Il expose "
        "une API REST servant d'interface entre l'application mobile et la base "
        "de donnees SQL Server de l'entreprise (instance SERVEURX3, instance X3V11)."
    )
    add_bullet_list(doc, [
        "Framework : Spring Boot 3.4.4 avec Spring Data JPA",
        "Base de donnees : SQL Server (via jTDS JDBC Driver 1.3.1)",
        "Port du serveur : 8080",
        "Instance : SERVEURX3 / X3V11",
        "Schema de base : MALLZELLIJ",
        "ORM : Hibernate (ddl-auto = none, schema existant)",
    ])

    doc.add_paragraph("3.3 Frontend - Application Android", style="Heading 2")
    doc.add_paragraph(
        "L'application Android est developpee en Java, ciblant l'API level 36 "
        "(Android 14) avec un minimum de API 24 (Android 7.0)."
    )
    add_bullet_list(doc, [
        "Langage : Java",
        "SDK minimum : API 24 (Android 7.0)",
        "SDK cible : API 36 (Android 14)",
        "Bibliotheques reseau : Retrofit 2 + OkHttp (avec logging)",
        "Scan de codes-barres : ZXing Embedded",
        "Design : Material Design 3 (Google Material)",
        "Serialisation JSON : Gson",
    ])

    doc.add_paragraph("3.4 Diagramme d'architecture", style="Heading 2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "+------------------+      HTTP/REST      +------------------+      JDBC      +------------------+\n"
        "|  Application     | <-----------------> |  Backend         | <------------> |  SQL Server      |\n"
        "|  Android         |   JSON (Gson)        |  Spring Boot     |   jTDS Driver  |  Sage X3         |\n"
        "|  (Retrofit 2)    |                      |  (Port 8080)     |                |  (SERVEURX3)     |\n"
        "+------------------+                      +------------------+                +------------------+\n"
        "                                                        |\n"
        "                                                  VMware VM\n"
        "                                                  (24h/24)"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    doc.add_paragraph()


def add_api_endpoints(doc):
    doc.add_paragraph("3.5 Endpoints API", style="Heading 2")
    add_table(doc,
        ["Methode", "Endpoint", "Description"],
        [
            ["POST", "/api/auth/login", "Authentification utilisateur"],
            ["GET", "/api/articles", "Liste de tous les articles"],
            ["GET", "/api/articles/search?q=", "Recherche d'articles par nom"],
            ["GET", "/api/articles/barcode/{ean}", "Recherche par code QR (ITMREF_0)"],
            ["GET", "/api/articles/{rowid}/stocks", "Stock detaille d'un article par site"],
            ["GET", "/api/inventory", "Liste des inventaires"],
            ["POST", "/api/inventory", "Creer un nouvel enregistrement d'inventaire"],
        ]
    )
    doc.add_paragraph()


def add_implementation(doc):
    doc.add_paragraph("Implementation - Les 17 Defis", style="Heading 1")

    doc.add_paragraph(
        "Le developpement a suivi un processus iteratif et incrementel, chaque "
        "challenge (CH) correspondant a une etape concrete du projet. Voici le "
        "detail de chaque defi releve."
    )

    # CH1
    doc.add_paragraph("4.1 CH1 - Choix des technologies et mise en place", style="Heading 2")
    doc.add_paragraph(
        "La premiere etape a consiste a choisir les technologies adaptees au projet : "
        "SQL Server pour la base de donnees (compatibilite avec le systeme Sage X3 "
        "de l'entreprise), Spring Boot pour le backend REST, et Android Studio pour "
        "le developpement de l'application mobile. L'environnement de developpement "
        "a ete configure : SDK Android, JDK 17, Gradle, et connexion a la base SQL "
        "Server existante via le driver jTDS."
    )

    # CH2
    doc.add_paragraph("4.2 CH2 - Interface utilisateur : login et register", style="Heading 2")
    doc.add_paragraph(
        "L'interface de connexion a ete developpee avec les composants Material "
        "Design (TextInputLayout, MaterialButton). Le formulaire de login verifie "
        "les identifiants aupres du backend via l'endpoint POST /api/auth/login. "
        "Les champs sont valides cote client avant l'envoi (champs obligatoires). "
        "En cas d'erreur, un message explicite est affiche. La classe UserSession "
        "(pattern Singleton) gere l'etat de connexion."
    )

    # CH3
    doc.add_paragraph("4.3 CH3 - Affichage des articles et des stocks par site", style="Heading 2")
    doc.add_paragraph(
        "Apres connexion, l'application affiche la liste des articles. Cinq tables "
        "sont impliquees : ITM_MASTER pour les articles, ITM_MVT pour les mouvements "
        "de stock, les sites (FBC, FMD) identifies dans les donnees. Le stock total "
        "d'un article est calcule en sommant les quantites PHYALL_0 de tous les sites. "
        "Lorsqu'un article est selectionne, un ecran detaille affiche le stock pour "
        "chaque site avec la quantite disponible et la quantite a louer."
    )

    # CH4
    doc.add_paragraph("4.4 CH4 - Deploiement sur VM et configuration reseau", style="Heading 2")
    doc.add_paragraph(
        "Le backend a ete deploie sur une machine virtuelle VMware fonctionnant 24h/24 "
        "afin de garantir la disponibilite du serveur. Un defi majeur a ete la "
        "connexion de l'appareil physique Android au serveur dans la VM. L'adresse IP "
        "de Retrofit a ete configuree (192.168.1.53:8080). Comme la connexion directe "
        "physique-VM etait impossible, une Android VM a ete creee avec ADB active et "
        "connectee via IP entre les VMs pour tester le projet."
    )

    # CH5
    doc.add_paragraph("4.5 CH5 - Connexion a la base de donnees de l'entreprise", style="Heading 2")
    doc.add_paragraph(
        "L'un des plus grands defis a ete de se connecter a la base de donnees "
        "existante de l'entreprise, qui utilise le systeme Sage X3. La base SQL "
        "Server contient des tables avec des conventions de nommage specifiques "
        "(champs suffixes par _0, tables prefixees par Y). Les annotations JPA "
        "@Column(name=...) ont ete utilisees pour mapper correctement les entites "
        "Java aux tables existantes sans modifier la base de donnees."
    )

    # CH6
    doc.add_paragraph("4.6 CH6 - Creation d'une base dediee et synchronisation", style="Heading 2")
    doc.add_paragraph(
        "Pour eviter de modifier la base Sage X3 existante, une nouvelle base de "
        "donnees a ete creee dans la meme instance SQL Server (MALLZELLIJ). Les "
        "tables necessaires ont ete importees depuis la base Sage X3. Un mecanisme "
        "de synchronisation a ete mis en place pour que les nouvelles donnees "
        "ajoutees aux tables originales soient automatiquement disponibles dans "
        "la base MALLZELLIJ."
    )

    # CH7
    doc.add_paragraph("4.7 CH7 - Scan de codes QR", style="Heading 2")
    doc.add_paragraph(
        "La fonctionnalite de scan de codes QR a ete integree via la bibliotheque "
        "ZXing Embedded. Un bouton camera a ete ajoute dans la barre de recherche. "
        "Lors du scan, la reference ITMREF_0 est extraite du code QR, puis envoyee "
        "au backend via l'endpoint /api/articles/barcode/{ean}. L'application affiche "
        "alors directement l'article, les sites et les quantites en stock, sans "
        "recherche manuelle. Un bip sonore confirme la detection du code."
    )

    # CH8
    doc.add_paragraph("4.8 CH8 - Deploiement sur serveur Windows", style="Heading 2")
    doc.add_paragraph(
        "Pour des raisons de securite, la VM Windows contenant le backend a ete "
        "copiee vers le serveur Windows de l'entreprise. L'application est accessible "
        "uniquement en reseau local (pas d'acces externe), seuls les appareils sur "
        "le meme reseau peuvent se connecter au serveur. Cette configuration "
        "garantit la securite des donnees de l'entreprise."
    )

    # CH9
    doc.add_paragraph("4.9 CH9 - Gestion des roles utilisateurs", style="Heading 2")
    doc.add_paragraph(
        "La table YMOBILE_USER a ete enrichie avec un champ de role (YROLE_0). "
        "Deux roles ont ete definis : admin et user1. Le role determine les "
        "fonctionnalites accessibles dans l'application. La verification se fait "
        "cote backend via la reponse d'authentification et cote frontend via la "
        "classe UserSession.getInstance().isAdmin()."
    )

    # CH10
    doc.add_paragraph("4.10 CH10 - Formulaire d'inventaire", style="Heading 2")
    doc.add_paragraph(
        "Un nouvel ecran d'inventaire a ete cree avec 4 types de champs : "
        "l'article (par scan QR code), les palettes, les cartons et les metres "
        "carres. Les donnees sont stockees dans la table YINV SQL Server. "
        "L'insertion se fait via l'endpoint POST /api/inventory. Un UUID est "
        "genere automatiquement pour chaque enregistrement."
    )

    # CH11
    doc.add_paragraph("4.11 CH11 - Sidebar de navigation", style="Heading 2")
    doc.add_paragraph(
        "Un menu lateral (sidebar) a ete ajoute avec NavigationView et "
        "DrawerLayout pour naviguer entre les ecrans : Articles, Inventaire "
        "et Deconnexion. Le menu s'adapte dynamiquement selon le role de "
        "l'utilisateur connecte."
    )

    # CH12
    doc.add_paragraph("4.12 CH12 - Controle d'acces par role", style="Heading 2")
    doc.add_paragraph(
        "L'utilisateur avec le role admin voit dans le sidebar : Articles, "
        "Inventaire et Deconnexion. L'utilisateur avec le role user1 ne voit "
        "que : Articles et Deconnexion. L'element Inventaire est masque "
        "dynamiquement via navView.getMenu().findItem(R.id.nav_inventory).setVisible(false)."
    )

    # CH13
    doc.add_paragraph("4.13 CH13 - Dropdowns equipe, depot, zone", style="Heading 2")
    doc.add_paragraph(
        "Des champs de selection (AutoCompleteTextView) ont ete ajoutes pour "
        "l'equipe, le depot et la zone. L'utilisateur choisit parmi des options "
        "predefinies (Equipe 1/2/3, zones et depots configures). L'ecriture "
        "directe est desactivee (InputType.TYPE_NULL) pour forcer la selection "
        "dans la liste. Un click affiche le dropdown."
    )

    # CH14
    doc.add_paragraph("4.14 CH14 - Persistance des selections", style="Heading 2")
    doc.add_paragraph(
        "Les selections equipe, depot et zone restent fixees apres le "
        "rechargement de la page, tandis que les autres champs (article, "
        "palettes, cartons, metres carres) se vident automatiquement apres "
        "la soumission. Cela optimise le flux de travail quand l'operateur "
        "travaille dans la meme zone/equipe."
    )

    # CH15
    doc.add_paragraph("4.15 CH15 - Enregistrement dynamique avec informations completes", style="Heading 2")
    doc.add_paragraph(
        "L'etape suivante consiste a enregistrer chaque inventaire avec toutes "
        "les informations : l'utilisateur qui a effectue la saisie, l'equipe, "
        "la zone, le depot, l'article, les palettes, les cartons, les metres "
        "carres, et la date de l'enregistrement. Cela permet un suivi complet "
        "des operations d'inventaire."
    )

    # CH16
    doc.add_paragraph("4.16 CH16 - Table YINV et dialog de confirmation", style="Heading 2")
    doc.add_paragraph(
        "La table YINV a ete creee dans SQL Server avec les colonnes necessaires. "
        "Le flux de soumission comprend : la validation des champs obligatoires, "
        "l'affichage d'un dialog de confirmation avec un tableau recapitulatif "
        "(equipe, depot, zone, article, palettes, cartons, metres carres, "
        "operateur), et deux boutons : Confirmer (insertion en base + reinitialisation "
        "du formulaire) et Editer (retour au formulaire pour correction)."
    )

    add_table(doc,
        ["Champ", "Type", "Description"],
        [
            ["Equipe", "Dropdown", "Equipe responsable de l'inventaire"],
            ["Depot", "Dropdown", "Depot concerne"],
            ["Zone", "Dropdown", "Zone de stockage"],
            ["Article", "Texte + QR", "Reference de l'article (scan possible)"],
            ["Palettes", "Numerique", "Nombre de palettes comptees"],
            ["Cartons", "Numerique", "Nombre de cartons comptes"],
            ["Metres carres", "Numerique", "Surface en m2"],
        ]
    )
    doc.add_paragraph()

    # CH17
    doc.add_paragraph("4.17 CH17 - Redesign et ajout du prix", style="Heading 2")
    doc.add_paragraph(
        "La derniere etape a consiste a redesigner l'interface des articles "
        "pour ameliorer l'ergonomie et ajouter l'affichage du prix unitaire. "
        "Le prix est recupere depuis la table SPRC_LIST (liste T11, code "
        "SPL26-0001) et affiche pour chaque article dans la vue detaillee "
        "du stock."
    )


def add_technologies(doc):
    doc.add_paragraph("Technologies Utilisees", style="Heading 1")

    doc.add_paragraph("5.1 Backend", style="Heading 2")
    add_table(doc,
        ["Technologie", "Version", "Role"],
        [
            ["Java", "17", "Langage de programmation backend"],
            ["Spring Boot", "3.4.4", "Framework web et REST API"],
            ["Spring Data JPA", "3.4.4", "Acces aux donnees via ORM"],
            ["Hibernate", "6.x", "ORM pour mapping objet-relationnel"],
            ["jTDS", "1.3.1", "Driver JDBC pour SQL Server"],
            ["Spring Security Crypto", "6.x", "Hachage des mots de passe"],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph("5.2 Frontend", style="Heading 2")
    add_table(doc,
        ["Technologie", "Version", "Role"],
        [
            ["Android SDK", "API 24-36", "Plateforme mobile"],
            ["Java", "11", "Langage de programmation mobile"],
            ["Retrofit", "2.x", "Client HTTP pour appels API"],
            ["OkHttp", "4.x", "Client reseau avec logging"],
            ["Gson", "2.x", "Serialisation JSON"],
            ["ZXing Embedded", "latest", "Lecture de codes-barres/QR"],
            ["Material Design", "3.x", "Composants UI"],
            ["RecyclerView", "latest", "Listes performantes"],
        ]
    )
    doc.add_paragraph()

    doc.add_paragraph("5.3 Outils de developpement", style="Heading 2")
    add_bullet_list(doc, [
        "Android Studio - IDE de developpement Android",
        "IntelliJ IDEA - IDE pour le backend Spring Boot",
        "Gradle 9.3.1 - Systeme de build",
        "VMware - Machine virtuelle pour deploiement serveur",
        "SQL Server Management Studio - Gestion de la base de donnees",
        "Git - Gestion de version",
    ])


def add_difficultes(doc):
    doc.add_paragraph("Difficultes Rencontrees et Solutions", style="Heading 1")

    doc.add_paragraph(
        "6.1 Connexion physique-VM : impossibilite de connecter un appareil "
        "physique Android a la VM VMware. Solution : creation d'une Android VM "
        "avec ADB active, connexion via IP entre les VMs.",
        style="Heading 2",
    )

    doc.add_paragraph(
        "6.2 Integration Sage X3 : la base de donnees de l'entreprise utilisait "
        "des conventions de nommage specifiques. Solution : annotations JPA "
        "@Column(name=...) pour mapper sans modifier la base existante.",
        style="Heading 2",
    )

    doc.add_paragraph(
        "6.3 Synchronisation des donnees : creation d'une base dediee MALLZELLIJ "
        "dans la meme instance SQL Server avec mecanisme de synchronisation.",
        style="Heading 2",
    )

    doc.add_paragraph(
        "6.4 Configuration reseau : plusieurs iterations pour trouver l'adresse "
        "IP correcte (192.168.1.53:8080). OkHttp logging facilite le diagnostic.",
        style="Heading 2",
    )

    doc.add_paragraph(
        "6.5 Securite : deploiement sur serveur local uniquement, pas d'acces "
        "externe au reseau de l'entreprise.",
        style="Heading 2",
    )


def add_bilan(doc):
    doc.add_paragraph("Bilan du Stage", style="Heading 1")

    doc.add_paragraph("7.1 Bilan technique", style="Heading 2")
    doc.add_paragraph(
        "Ce stage m'a permis de mettre en pratique mes connaissances en "
        "developpement mobile et backend. J'ai acquis une experience concrete dans :"
    )
    add_bullet_list(doc, [
        "Le developpement d'applications Android avec Java et Material Design",
        "La conception et l'implementation d'API REST avec Spring Boot",
        "L'integration d'une base de donnees existante (Sage X3) via JPA/Hibernate",
        "La gestion de la communication reseau (Retrofit, OkHttp, configuration IP)",
        "L'integration de fonctionnalites materielles (scan de codes QR)",
        "Le deploiement sur VM VMware pour un service 24h/24",
        "La gestion des droits d'acces par role (admin/user1)",
        "Le travail en equite et la collaboration avec les services metiers",
    ])

    doc.add_paragraph("7.2 Bilan personnel", style="Heading 2")
    doc.add_paragraph(
        "Ce stage a ete une experience extremement enrichissante sur le plan "
        "personnel et professionnel. J'ai appris a travailler en autonomie tout "
        "en sollicitant les conseils de mon encadrant Driss Rokh. Les 17 defis "
        "releves m'ont appris la persistance, la resolution de problemes et "
        "l'adaptation aux contraintes techniques reelles d'une entreprise."
    )

    doc.add_paragraph("7.3 Perspectives d'amelioration", style="Heading 2")
    add_numbered_list(doc, [
        "Implementer un systeme d'authentification par jeton JWT",
        "Ajouter un mode hors ligne avec synchronisation differee",
        "Integrer des notifications push pour les alertes de stock",
        "Developper un tableau de bord analytique pour le suivi des inventaires",
        "Automatiser completement la synchronisation avec Sage X3",
    ])


def add_conclusion(doc):
    doc.add_paragraph("Conclusion", style="Heading 1")
    doc.add_paragraph(
        "Ce stage de six semaines chez Mallzellij a ete l'occasion de concevoir et "
        "developper une application Android complete de gestion d'inventaire. Le "
        "projet a suivi un processus iteratif de 17 defis techniques, couvrant "
        "toutes les etapes du developpement logiciel : du choix des technologies "
        "au deploiement sur serveur."
    )
    doc.add_paragraph(
        "L'application developpee offre une solution moderne et efficace pour la "
        "gestion des stocks, avec des fonctionnalites de scan de codes QR, de "
        "recherche rapide, de formulaire d'inventaire intuitif et de gestion "
        "des roles utilisateurs. L'integration avec la base de donnees Sage X3 "
        "de l'entreprise a ete un defi particulierement instructif."
    )
    doc.add_paragraph(
        "Je tiens a remercier mon encadrant Driss Rokh pour sa disponibilite et "
        "ses conseils tout au long de ce stage, ainsi que l'equipe Mallzellij pour "
        "son accueil chaleureux et sa confiance. Ce stage a confirme ma passion "
        "pour le developpement logiciel et m'a prepare aux defis professionnels a venir."
    )


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    add_cover(doc)
    doc.add_page_break()

    doc.add_paragraph("Sommaire", style="Heading 1")
    add_toc(doc)
    doc.add_page_break()

    add_general_introduction(doc)
    doc.add_page_break()

    add_presentation_entreprise(doc)
    doc.add_page_break()

    add_conception(doc)
    doc.add_page_break()

    add_architecture(doc)
    doc.add_page_break()

    add_api_endpoints(doc)
    doc.add_page_break()

    add_implementation(doc)
    doc.add_page_break()

    add_technologies(doc)
    doc.add_page_break()

    add_difficultes(doc)
    doc.add_page_break()

    add_bilan(doc)
    doc.add_page_break()

    add_conclusion(doc)

    setup_header_footer(doc)

    output_path = os.path.join(os.path.dirname(__file__), "Rapport_Stage_Abdelkarim_Erraji_v2.docx")
    doc.save(output_path)
    print(f"Rapport genere avec succes : {output_path}")


if __name__ == "__main__":
    main()
